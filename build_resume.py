from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

# --- Page margins ---
for section in doc.sections:
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)

# --- Default style ---
doc.styles['Normal'].font.name = 'Calibri'
doc.styles['Normal'].font.size = Pt(11)


def _set_spacing(p, before=0, after=2):
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)


def add_name(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(22)
    run.font.name = 'Calibri'
    _set_spacing(p, before=0, after=2)


def add_contact(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(10)
    run.font.name = 'Calibri'
    _set_spacing(p, before=0, after=8)


def add_section_header(text):
    p = doc.add_paragraph()
    run = p.add_run(text.upper())
    run.bold = True
    run.font.size = Pt(11)
    run.font.name = 'Calibri'
    # Bottom border line
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '000000')
    pBdr.append(bottom)
    pPr.append(pBdr)
    _set_spacing(p, before=10, after=3)


def add_body(text, size=10.5):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.name = 'Calibri'
    _set_spacing(p, before=0, after=2)


def add_job_header(title, company_loc, dates):
    p = doc.add_paragraph()
    r1 = p.add_run(f'{title}  |  {company_loc}')
    r1.bold = True
    r1.font.size = Pt(11)
    r1.font.name = 'Calibri'
    _set_spacing(p, before=8, after=0)
    p2 = doc.add_paragraph()
    r2 = p2.add_run(dates)
    r2.italic = True
    r2.font.size = Pt(10)
    r2.font.name = 'Calibri'
    _set_spacing(p2, before=0, after=2)


def add_bullet(text, size=10.5):
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.name = 'Calibri'
    p.paragraph_format.left_indent = Inches(0.25)
    _set_spacing(p, before=1, after=1)


def add_skills_line(label, skills):
    p = doc.add_paragraph()
    rl = p.add_run(f'{label}: ')
    rl.bold = True
    rl.font.size = Pt(10.5)
    rl.font.name = 'Calibri'
    rs = p.add_run(skills)
    rs.font.size = Pt(10.5)
    rs.font.name = 'Calibri'
    _set_spacing(p, before=1, after=2)


# ════════════════════════════════════════════════════════
#  HEADER
# ════════════════════════════════════════════════════════
add_name('Md Mehedi Hasan')
add_contact(
    'United States (Remote)  |  mhasan.kme@gmail.com  |  (248) 861-9788  |  '
    'linkedin.com/in/mehedihasan91'
)

# ════════════════════════════════════════════════════════
#  PROFESSIONAL SUMMARY
# ════════════════════════════════════════════════════════
add_section_header('Professional Summary')
add_body(
    'Senior Software Development Engineer in Test (SDET) with 5+ years of experience building and scaling '
    'Python-based test automation platforms and CI/CD pipelines used by multiple engineering teams. '
    'Proven track record delivering measurable improvements in test reliability, release velocity, and '
    'developer productivity through robust test infrastructure and DevOps-driven quality practices. '
    'Deep expertise in framework ownership, automated regression testing, and cross-functional collaboration '
    'in fast-paced Agile/Scrum environments. U.S. citizen, remote-ready.'
)

# ════════════════════════════════════════════════════════
#  KEY ACHIEVEMENTS
# ════════════════════════════════════════════════════════
add_section_header('Key Achievements')
for bullet in [
    'Designed and deployed a Python-based test automation platform supporting large-scale system and '
    'integration testing across 5+ engineering teams, resulting in improved test reliability and '
    'faster regression cycles.',

    'Optimized CI/CD pipelines and enabled parallel test execution using Jenkins and GitHub Actions, '
    'accelerating validation and release timelines and improving feedback speed for development teams.',

    'Automated environment provisioning and deployment workflows using Docker and Ansible, eliminating '
    'manual setup steps and increasing consistency across test environments.',

    'Served as framework owner and primary technical point of contact for shared automation tooling, '
    'driving adoption, troubleshooting, and continuous improvement across multiple engineering teams.',
]:
    add_bullet(bullet)

# ════════════════════════════════════════════════════════
#  TECHNICAL SKILLS
# ════════════════════════════════════════════════════════
add_section_header('Technical Skills')
add_skills_line('Languages & Scripting',      'Python, Java, C/C++, MATLAB')
add_skills_line('Test Automation & Frameworks','Robot Framework, Selenium, Appium, TestNG, JUnit, Cucumber, Postman')
add_skills_line('CI/CD & DevOps',             'Jenkins, GitHub Actions, Docker, Kubernetes, Terraform, Ansible, Puppet')
add_skills_line('Cloud & Monitoring',         'AWS, Azure, Grafana')
add_skills_line('Version Control & Tools',    'Git, Jira')
add_skills_line('Methodologies',              'Agile, Scrum, Test-Driven Development (TDD), DFSS Green & Black Belt')

# ════════════════════════════════════════════════════════
#  PROFESSIONAL EXPERIENCE
# ════════════════════════════════════════════════════════
add_section_header('Professional Experience')

# -- GM Warren --
add_job_header(
    'Test Automation Framework Engineer',
    'General Motors — Warren, MI',
    'April 2024 – Present'
)
for b in [
    'Own and scale a Python-based test automation platform supporting system and integration testing '
    'across multiple engineering teams, improving test reliability and reducing flaky test rates.',

    'Design and maintain Jenkins and GitHub Actions CI/CD pipelines enabling automated test execution, '
    'real-time reporting, and release validation with reduced manual intervention.',

    'Partner with developers and system engineers to define test strategies and increase automated test '
    'coverage across critical system components.',

    'Author and maintain technical documentation for automation frameworks and test environments, '
    'accelerating onboarding time for new team members.',

    'Lead continuous improvement initiatives improving test execution speed through parallel execution '
    'and pipeline optimization.',
]:
    add_bullet(b)

# -- GM Milford --
add_job_header(
    'Test Automation Engineer',
    'General Motors — Milford, MI',
    'September 2021 – March 2024'    # FIXED: was September 2020
)
for b in [
    'Built and extended Python-based automated test suites for system and integration testing across '
    'multiple product lines, contributing to reduced regression cycle time and improved release confidence.',

    'Implemented and maintained Jenkins CI/CD pipelines for scheduled, nightly, and on-demand test '
    'execution with automated reporting, enabling faster feedback loops for development teams.',

    'Translated system requirements into effective test strategies and automated test coverage in close '
    'collaboration with developers and system engineers.',

    'Created technical documentation and onboarding materials enabling rapid adoption of the automation '
    'framework across engineering teams.',

    'Diagnosed and resolved test failures and infrastructure issues, improving execution stability and '
    'reducing test flakiness over time.',
]:
    add_bullet(b)

# -- Stellantis --
add_job_header(
    'Software Modeling Engineer',
    'Stellantis North America — Auburn Hills, MI',
    'August 2017 – September 2021'
)
for b in [
    'Designed and implemented embedded software components for digital user interfaces and display '
    'systems in automotive platforms.',

    'Developed model-based software using MATLAB/Simulink and validated software quality through '
    'automated and manual testing processes.',

    'Collaborated cross-functionally with product, design, and quality teams to deliver reliable, '
    'user-facing software features aligned with functional requirements.',

    'Performed functional validation, debugging, and performance analysis to ensure software stability '
    'and quality prior to production release.',
]:
    add_bullet(b)

# -- Georgia Southern --
add_job_header(
    'Graduate Research Assistant',
    'Georgia Southern University — Statesboro, GA',
    'August 2015 – May 2017'
)
for b in [
    'Conducted research in power electronics and distributed energy systems.',
    'Developed simulation models and optimization algorithms using MATLAB and Simulink.',
    'Supported undergraduate instruction and laboratory experimentation.',
]:
    add_bullet(b)

# ════════════════════════════════════════════════════════
#  EDUCATION
# ════════════════════════════════════════════════════════
add_section_header('Education')
add_body('Master of Science (M.S.) in Electrical Engineering  —  Georgia Southern University, Statesboro, GA  |  May 2017')
add_body('Bachelor of Science (B.S.) in Mechanical Engineering  —  Khulna University of Engineering & Technology (KUET), Bangladesh  |  November 2013')

# ════════════════════════════════════════════════════════
#  VOLUNTEER WORK
# ════════════════════════════════════════════════════════
add_section_header('Volunteer Work')
add_bullet('Volunteer, GM Cares and Give Aways')
add_bullet('Former Global Ambassador, Georgia Southern University, USA')

# ════════════════════════════════════════════════════════
#  SAVE
# ════════════════════════════════════════════════════════
output = 'Md_Mehedi_Hasan_Resume_ATS.docx'
doc.save(output)
print(f'Saved: {output}')
